#!/usr/bin/env python 3.9
# -*- coding: utf-8 -*-
# @Time    : 2022/12/6 18:02 update
# @Author  : ZCG
# @File    : WordReplace.py
# @Software: PyCharm
# @Notice  :
#pip install pyopenssl --upgrade
#https://stackoverflow.com/questions/34779724/python-docx-replace-string-in-paragraph-while-keeping-style
from docx import Document
import os
from boto3.s3.transfer import S3Transfer
import boto3
import requests
import docx
import io
import tempfile

class Execute:
    '''
        Execute Paragraphs KeyWords Replace
        paragraph: docx paragraph
    '''

    def __init__(self, paragraph):
        self.paragraph = paragraph


    def p_replace(self, x:int, key:str, value:str):
        '''
        paragraph replace
        The reason why you do not replace the text in a paragraph directly is that it will cause the original format to
        change. Replacing the text in runs will not cause the original format to change
        :param x:       paragraph id
        :param key:     Keywords that need to be replaced
        :param value:   The replaced keywords
        :return:
        '''
        # Gets the coordinate index values of all the characters in this paragraph [{run_index , char_index}]
        p_maps = [{"run": y, "char": z} for y, run in enumerate(self.paragraph.runs) for z, char in enumerate(list(run.text))]
        # Handle the number of times key occurs in this paragraph, and record the starting position in the list.
        # Here, while self.text.find(key) >= 0, the {"ab":"abc"} term will enter an endless loop
        # Takes a single paragraph as an independent body and gets an index list of key positions within the paragraph, or if the paragraph contains multiple keys, there are multiple index values
        k_idx = [s for s in range(len(self.paragraph.text)) if self.paragraph.text.find(key, s, len(self.paragraph.text)) == s]
        for i, start_idx in enumerate(reversed(k_idx)):       # Reverse order iteration
            end_idx = start_idx + len(key)                    # The end position of the keyword in this paragraph
            k_maps = p_maps[start_idx:end_idx]                # Map Slice List A list of dictionaries for sections that contain keywords in a paragraph
            self.r_replace(k_maps, value)
            print(f"\t |Paragraph {x+1: >3}, object {i+1: >3} replaced successfully! | {key} ===> {value}")


    def r_replace(self, k_maps:list, value:str):
        '''
        :param k_maps: The list of indexed dictionaries containing keywords， e.g:[{"run":15, "char":3},{"run":15, "char":4},{"run":16, "char":0}]
        :param value:
        :return:
        Accept arguments, removing the characters in k_maps from back to front, leaving the first one to replace with value
        Note: Must be removed in reverse order, otherwise the list length change will cause IndedxError: string index out of range
        '''
        for i, position in enumerate(reversed(k_maps), start=1):
            y, z = position["run"], position["char"]
            run:object = self.paragraph.runs[y]         # "k_maps" may contain multiple run ids, which need to be separated
            # Pit: Instead of the replace() method, str is converted to list after a single word to prevent run.text from making an error in some cases (e.g., a single run contains a duplicate word)
            thisrun = list(run.text)
            if i < len(k_maps):
                thisrun.pop(z)          # Deleting a corresponding word
            if i == len(k_maps):        # The last iteration (first word), that is, the number of iterations is equal to the length of k_maps
                thisrun[z] = value      # Replace the word in the corresponding position with the new content
            run.text = ''.join(thisrun) # Recover



class WordReplace:
    '''
        file: Microsoft Office word file，only support .docx type file
    '''

    def __init__(self, file_url):
        #file_url = 'https://images-storage-bucket.s3.ap-southeast-1.amazonaws.com/upload/avatar/files/1684466053091-Dai%20hoi%20XIII.docx'

        # Download the file content from the S3 server
        response = requests.get(file_url)
        file_content = response.content
        #print(file_content)
        # Open the downloaded file using python-docx
        self.docx = docx.Document(io.BytesIO(file_content))

    def body_content(self, replace_dict:dict):
        print("\t☺Processing keywords in the body...")
        for key, value in replace_dict.items():
            for x, paragraph in enumerate(self.docx.paragraphs):
                #print(paragraph)
                Execute(paragraph).p_replace(x, key, value)
        print("\t |Body keywords in the text are replaced!")


    def body_tables(self,replace_dict:dict):
        print("\t☺Processing keywords in the body'tables...")
        for key, value in replace_dict.items():
            for table in self.docx.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for x, paragraph in enumerate(cell.paragraphs):
                            Execute(paragraph).p_replace(x, key, value)
        print("\t |Body'tables keywords in the text are replaced!")


    def header_content(self,replace_dict:dict):
        print("\t☺Processing keywords in the header'body ...")
        for key, value in replace_dict.items():
            for section in self.docx.sections:
                for x, paragraph in enumerate(section.header.paragraphs):
                    Execute(paragraph).p_replace(x, key, value)
        print("\t |Header'body keywords in the text are replaced!")


    def header_tables(self,replace_dict:dict):
        print("\t☺Processing keywords in the header'tables ...")
        for key, value in replace_dict.items():
            for section in self.docx.sections:
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for x, paragraph in enumerate(cell.paragraphs):
                                Execute(paragraph).p_replace(x, key, value)
        print("\t |Header'tables keywords in the text are replaced!")


    def footer_content(self, replace_dict:dict):
        print("\t☺Processing keywords in the footer'body ...")
        for key, value in replace_dict.items():
            for section in self.docx.sections:
                for x, paragraph in enumerate(section.footer.paragraphs):
                    Execute(paragraph).p_replace(x, key, value)
        print("\t |Footer'body keywords in the text are replaced!")


    def footer_tables(self, replace_dict:dict):
        print("\t☺Processing keywords in the footer'tables ...")
        for key, value in replace_dict.items():
            for section in self.docx.sections:
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for x, paragraph in enumerate(cell.paragraphs):
                                Execute(paragraph).p_replace(x, key, value)
        print("\t |Footer'tables keywords in the text are replaced!")

    def extract_text_after_last_slash(self,url):
        # Example usage
        # url = 'https://example.com/path/to/file.docx'
        # text = extract_text_after_last_slash(url)
        # print(text)  # Output: file.docx
        return url.rsplit('/', 1)[-1]
    

    def upload_file_to_s3(self,temp_file,file_url):
    # Create an S3 client
        S3_ACCESS_KEY = "AKIA4I73TQ6VHPIJUNET"
        S3_SECRET_KEY = "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"
        bucket_name = "images-storage-bucket"
        client = boto3.client('s3',
                        aws_access_key_id=S3_ACCESS_KEY,
                        aws_secret_access_key=S3_SECRET_KEY)
        transfer = S3Transfer(client)
        # Upload the file to S3
        fileName= self.extract_text_after_last_slash(file_url)

        key= 'upload/avatar/files/'+fileName[:-5] + '_translated.docx'
        try:
            transfer.upload_file(temp_file.name, bucket_name, key)
            print("File uploaded successfully.")
        except Exception as e:
            print("Error uploading file:", e)

    def save(self, filepath:str):
        '''
        :param filepath: File saving path
        :return:
        '''
        self.docx.save(filepath[:-5] + '_translated.docx')

    def save_document(self, document):
        '''
        :document: represents the document object"docx.document"
        :return: temp_file : Temporary file object docx (to upload to S3)
        '''
        doc_buffer = io.BytesIO()
        document.save(doc_buffer)
        doc_buffer.seek(0)

        #create a temporary file to save the content of the buffer
        temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        temp_file.write(doc_buffer.getvalue())
        temp_file.close()
        return temp_file
    



        



    @staticmethod
    def docx_list(dirPath):
        '''
        :param dirPath:
        :return: List of docx files in the current directory
        '''
        fileList = []
        for roots, dirs, files in os.walk(dirPath):
            for file in files:
                if file.endswith("docx") and file[0] != "~":  # Find the docx document and exclude temporary files
                    fileRoot = os.path.join(roots, file)
                    fileList.append(fileRoot)
        print("This directory finds a total of {0} related files!".format(len(fileList)))
        print(fileList)
        return fileList


    


'''
To use: Modify the values in replace dict and filedir
replace_dict ：key:to be replaced, value:new content
filedir ：Directory where docx files are stored. Subdirectories are supported
'''


# input section
# replace_dict = {
#     "Jonathan Yuen, your feeling of being disgusted of the wrong grammatical expression of the other people doing their comments about the story The Chapel DOES NOT HELP THEM! Instead of feeling disgusted, why don’t you specifically tell them the wrong you have noticed in a very friendly manner so they may know their wrong and learn from such mistake? Your words do speak the kind of attitude you do have. Change it! If not, control it":"Jonathan Yuen, cảm giác ghê tởm của bạn về cách diễn đạt ngữ pháp sai của những người khác khi nhận xét về câu chuyện Nhà nguyện KHÔNG GIÚP HỌ! Thay vì cảm thấy ghê tởm, tại sao bạn không nói cụ thể với họ những sai lầm mà bạn đã nhận thấy một cách rất thân thiện để họ có thể biết sai của họ và học hỏi từ sai lầm đó? Lời nói của bạn nói lên loại thái độ mà bạn có. Thay đổi nó! Nếu không, hãy kiểm soát nó",
#     "Aaa":"hehe",
#     }
# filedir = r"C:\Users\DELL\Desktop\New folder"

# def test(filename):
#     wordreplace = WordReplace(filename)
#     wordreplace.header_content(replace_dict)
#     wordreplace.header_tables(replace_dict)
#     wordreplace.body_content(replace_dict)
#     wordreplace.body_tables(replace_dict)
#     wordreplace.footer_content(replace_dict)
#     wordreplace.footer_tables(replace_dict)
#     wordreplace.save(filename)
#     print(f'\t☻The document processing is complete!\n')

# test(r"C:\Users\DELL\Desktop\New folder\vi.docx")


def export_file(sents,translated_sents,file_url):
    replace_dict = {}
    for i in range(len(sents)):
        replace_dict[sents[i]] = translated_sents[i]
    
    print(replace_dict)

    wordreplace = WordReplace(file_url)
    wordreplace.header_content(replace_dict)
    wordreplace.header_tables(replace_dict)
    wordreplace.body_content(replace_dict)
    wordreplace.body_tables(replace_dict)
    wordreplace.footer_content(replace_dict)
    wordreplace.footer_tables(replace_dict)
    temp_file = wordreplace.save_document(wordreplace.docx)
    #print(temp_file)
    wordreplace.upload_file_to_s3(temp_file,file_url)
    #wordreplace.save(file_url)
    print(f'\t☻The document processing is complete!\n')

    
sents=["1. Đại hội XIII","- Đại hội đại biểu toàn quốc lần thứ VIII và bước đầu thực hiện công cuộc đẩy mạnh công nghiệp hoá, hiện đại hoá 1996-2001"]
translated_sents=["1. XIII Congress","- The Eighth National Congress of Deputies and the initial implementation of the promotion of industrialization and modernization 1996-2001"]

replace_dict = {}
for i in range(len(sents)):
    replace_dict[sents[i]] = translated_sents[i]
        
#remove pair {" ":" "} in replace_dict
for key in list(replace_dict.keys()):
    if key == "":
        del replace_dict[key]
    if replace_dict[key] == " ":
        del replace_dict[key]


print(replace_dict)
#file_url = 'https://images-storage-bucket.s3.ap-southeast-1.amazonaws.com/upload/avatar/files/1684466053091-Dai%20hoi%20XIII.docx'

#test
#export_file(sents,translated_sents,file_url)

#pip install pyopenssl --upgrade
# from boto3.s3.transfer import S3Transfer
# import boto3
# #have all the variables populated which are required below

# import boto3
# def upload_file_to_s3(file_path, bucket_name):
#     # Create an S3 client
#     S3_ACCESS_KEY = "AKIA4I73TQ6VHPIJUNET"
#     S3_SECRET_KEY = "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"
#     client = boto3.client('s3',
#                       aws_access_key_id=S3_ACCESS_KEY,
#                       aws_secret_access_key=S3_SECRET_KEY)
#     transfer = S3Transfer(client)
#     # Upload the file to S3
#     key= "upload/avatar/files/vi_translated.docx"
#     try:
#         transfer.upload_file(file_path, bucket_name, key)
#         print("File uploaded successfully.")
#     except Exception as e:
#         print("Error uploading file:", e)


#S3_URI = https://images-storage-bucket.s3.ap-southeast-1.amazonaws.com/
# bucket_name = "images-storage-bucket"
# file_path = r"C:\Users\DELL\Desktop\New folder\vi_translated.docx"
# upload_file_to_s3(file_path, bucket_name)



# Call processing section
# for i, file in enumerate(WordReplace.docx_list(filedir),start=1):
#     print(i)
#     print("ao the nho")
#     print(f"{i}、Processing file:{file}")
#     wordreplace = WordReplace(file)
#     wordreplace.header_content(replace_dict)
#     wordreplace.header_tables(replace_dict)
#     wordreplace.body_content(replace_dict)
#     wordreplace.body_tables(replace_dict)
#     wordreplace.footer_content(replace_dict)
#     wordreplace.footer_tables(replace_dict)
#     wordreplace.save(file)
#     print(f'\t☻The document processing is complete!\n')


