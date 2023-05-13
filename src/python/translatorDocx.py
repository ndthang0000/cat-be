# pip install python-docx
# pip install deep-translator
import docx
from deep_translator import GoogleTranslator
from underthesea import word_tokenize
from underthesea import sent_tokenize
import sys
# deep_translator dung gg API

# ---------------------------------------------------------
# function 1: function tách câu từ file : input là tên file , output là array đã các câu được tách


def sent_from_file(filename):
    docx_file = docx.Document(r"uploads/"+filename)
    sent = []
    for paragraph in docx_file.paragraphs:
        text = paragraph.text
        sent_Vi = sent_tokenize(text)
        sent = sent + sent_Vi
    return sent


# ---------------------------------------------------------
# function 2: function dịch 1 câu: input là 1 câu kem theo ngon ngu can dich; output là 1 câu đã đc dịch
def translate_sent(sent, language):
    translated_text = GoogleTranslator(
        source='auto', target=language).translate(sent)
    return translated_text


# ---------------------------------------------------------
# function 3: function dịch nhiều câu: input là array các câu chưa dịch va ngon ngu can dich, output là mảng các câu đc dịch
def translate_sents(sents, language):
    translated_sents = []
    for sent in sents:
        translated_sent = translate_sent(sent, language)
        translated_sents.append(translated_sent)
    return translated_sents


# test
'''
filename = "vietnamtext.docx"
language="en"
#sent = sent_from_file(filename)
#print(sent)
print("-------------------------------------------------")
sent='Cuộc phản công mùa xuân thành công sẽ quyết định vận mệnh cuộc chiến'
print(sent)
translated_sent = translate_sents(sent,language)
print(translated_sent)
'''

def main():
    if sys.argv[1] == "sent_from_file":
        sent = sent_from_file(sys.argv[2])
        for i in sent:
            print(i.encode('utf-8'))
    elif sys.argv[1] == "translate_sent":
        translated_sent = translate_sent(sys.argv[2], sys.argv[3])
        print(translated_sent.encode('utf-8'))
        # print(translated_sent.encode('utf-8'))
    elif sys.argv[1] == "translate_sents":
        print(sys.argv[2])
        translated_sents = translate_sents(sys.argv[2], sys.argv[3])
        print(translated_sents)
    else:
        print("error")

main()


# ---------------------------------------------------------
# de day neu can xem lai


# test Translate
# ---------------------------------------------------

#translated = GoogleTranslator(source='auto', target='vi').translate("keep it up, you are awesome")
# print(translated)

# default return type is a list
# langs_list = GoogleTranslator().get_supported_languages()  # output: [arabic, french, english etc...]
# print(langs_list)
# langs_dict = GoogleTranslator().get_supported_languages(as_dict=True)  # output: {arabic: ar, french: fr, english:en etc...}
# print(langs_dict)

'''
#-----------------------------------------------------------------------------------------
# Mở file Docx gốc và file Docx mới : anylanguage! =)))
docx_file = docx.Document(r"uploads\1683339435548-MSSV.docx")
new_docx_file = docx.Document()
# Dịch văn bản từ file Docx gốc sang tiếng Việt và lưu vào file Docx mới
all_paras = docx_file.paragraphs
print(len(all_paras))
for paragraph in docx_file.paragraphs:
    #print(paragraph.text)
    #print("-----------------------")
    text = paragraph.text
    translated_text = GoogleTranslator(source='auto', target='en').translate(text)
    new_docx_file.add_paragraph(translated_text)

# Lưu trữ đối tượng Docx mới vào file Docx mới
new_docx_file.save('uploads/new.docx')
'''
