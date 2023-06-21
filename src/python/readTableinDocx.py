import docx
import pandas as pd
import numpy as np


#Parsing table from docx using Python
doc = docx.Document(r"uploads/"+'vi.docx')

df = pd.DataFrame()

tables = doc.tables[0]

##Getting the original data from the document to a list
ls =[]
for row in tables.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            ls.append(paragraph.text)

print(ls)

def Doctable(ls, row, column):
    df = pd.DataFrame(np.array(ls).reshape(row,column))  #reshape to the table shape
    new = docx.Document()
    word_table =new.add_table(rows = row, cols = column)
    for x in range(0,row,1):
        for y in range(0,column,1):
            cell = word_table.cell(x,y)
            cell.text = df.iloc[x,y]


    return new, df

import requests
import docx
import io

# Specify the URL link to the DOCX file on the S3 server
file_url = 'https://images-storage-bucket.s3.ap-southeast-1.amazonaws.com/upload/avatar/files/1684466053091-Dai%20hoi%20XIII.docx'

# Download the file content from the S3 server
response = requests.get(file_url)
file_content = response.content
#print(file_content)
# Open the downloaded file using python-docx
document = docx.Document(io.BytesIO(file_content))

print(io.BytesIO(file_content))
# Access and manipulate the content of the document
for paragraph in document.paragraphs:
    print(paragraph.text)

#test 1: checked

#upload file to s3
import boto3
import os
import uuid
from botocore.exceptions import ClientError
